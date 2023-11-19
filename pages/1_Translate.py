import streamlit as st
import db
import openai
from docx import Document
import base64
import io  # Don't forget to import the 'io' module
openai.api_key = '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'
textToReturn = 'test'
def translate_text(input_text, source_language, target_language):
    # Formulate a prompt for translation
    prompt = f"Translate the following text from {source_language} to {target_language}: '{input_text}'"

    # Specify the OpenAI GPT-3 engine and provide the prompt
    response = openai.Completion.create(
        engine="text-davinci-003",  # You can choose a different engine
        prompt=prompt,
        max_tokens=150,  # Adjust the desired length of the response
        n=1,  # Number of completions to generate
        stop=None,  # You can specify stop sequences to limit the response
    )
    
    # Extract and return the generated translation
    translated_text = response.choices[0].text.strip()
    return translated_text
def generate_word_document(content):
    doc = Document()
    doc.add_heading('Streamlit Generated Word Document', 0)

    # Add content to the document
    doc.add_paragraph(content)

    # Save the document
    doc.save('streamlit_generated_document.docx')

    st.success('Word document generated successfully!')
if "user" not in st.experimental_get_query_params():
    st.experimental_set_query_params(user="no")
if st.experimental_get_query_params()["user"][0] == "no":
    st.title("Login to see this page")

else:
    # Example usage
    st.title("Trouver your Expressions!")
    source_language = st.text_input("Enter source_language:")
    target_language = st.text_input("Enter target_language:")
    input_text = st.text_input("Enter desired expressions:")
    # Generate and print the translation
    


    # Create a button
    if st.button('Click me'):
    # Code to execute when the button is clicked
        translation = translate_text(input_text, source_language, target_language)
        st.write("Translated text:",translation)
        textToReturn = input_text + '\n' + translation
    

def generate_word_document(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    # Save the document to a BytesIO object
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes

st.title("Word Document Generator")

# Get user input
document_title = st.text_input("Enter document title:")
document_content = st.text_area("Enter document content:")

document_content=document_content.replace('.','.\n')
document_content=document_content.replace('?','?\n')
# Button to generate the Word document
if st.button("Generate Word Document"):
    doc_bytes = generate_word_document(document_title, document_content)

    # Create a downloadable link
    st.download_button(
        label="Download Word Document",
        data=doc_bytes,
        file_name="generated_document.docx",
        key="word-document-download"
    )

with open("footer.html") as f:
    foot = f.read()

with open("styles.css") as f:
    styling = f.read()

st.markdown(foot, unsafe_allow_html=True)
st.markdown(f"""<style>{styling}</style>""", unsafe_allow_html=True)